import os
import pandas as pd
from flask import Flask, render_template, request, send_file
from analyzer import process_excel_file, calculate_stats
from llm_engine import generate_marketing_summary
from validator import run_gold_standard_validation
from config import BATCH_SIZE # Import fixed configuration

# Import libraries for export
from bs4 import BeautifulSoup
from docx import Document

app = Flask(__name__)

# Configure paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
os.makedirs(DATA_DIR, exist_ok=True)

# Temporary file paths
TEMP_EXCEL_PATH = os.path.join(DATA_DIR, "latest_results.xlsx")
TEMP_SUMMARY_PATH = os.path.join(DATA_DIR, "latest_summary.html")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    # 1. Basic file checks
    if 'file' not in request.files: return "No file uploaded", 400
    file = request.files['file']
    if file.filename == '': return "No file selected", 400

    # 2. Save the uploaded file
    file_path = os.path.join(DATA_DIR, "uploaded_reviews.xlsx")
    file.save(file_path)
    
    # 3. Get parameters (Batch size retrieval is no longer needed; using constant directly)
    dual_mode = request.form.get('dual_mode') == 'on'
    
    print(f"Server: Start Analysis. Fixed Batch Size: {BATCH_SIZE}, Dual Mode: {dual_mode}")

    # 4. Processing workflow
    stats_note = ""
    consistency_score = None

    if dual_mode:
        # --- Dual Validation Mode (Repeatability) ---
        print(">>> Running Pass 1...")
        df_run1 = process_excel_file(file_path, batch_size=BATCH_SIZE)
        
        print(">>> Running Pass 2 (Verification)...")
        df_run2 = process_excel_file(file_path, batch_size=BATCH_SIZE)
        
        df_results, consistency_score = compare_and_merge_runs(df_run1, df_run2)
        stats_note = f"\n(Note: Analysis verified with dual-run consistency of {consistency_score:.2f}%.)"
    else:
        # --- Standard Mode ---
        df_results = process_excel_file(file_path, batch_size=BATCH_SIZE)

    # 5. Aggregate statistics
    stats_df = calculate_stats(df_results)
    
    # 6. Generate report
    summary_html = generate_marketing_summary(stats_df.to_string() + stats_note)

    # Save results for download
    try:
        df_results.to_excel(TEMP_EXCEL_PATH, index=False)
        with open(TEMP_SUMMARY_PATH, 'w', encoding='utf-8') as f:
            content = summary_html if summary_html else "No summary available."
            f.write(content)
    except Exception as e:
        print(f"Error saving temp files: {e}")

    return render_template(
        'dashboard.html',
        tables=[df_results.to_html(classes='table table-striped', index=False)],
        stats_table=stats_df.to_html(classes='table table-bordered'),
        summary=summary_html,
        consistency=consistency_score
    )

@app.route('/download_data')
def download_data():
    if os.path.exists(TEMP_EXCEL_PATH):
        return send_file(TEMP_EXCEL_PATH, as_attachment=True, download_name="review_analysis_data.xlsx")
    return "No data generated yet.", 404

@app.route('/download_report')
def download_report():
    if not os.path.exists(TEMP_SUMMARY_PATH):
        return "No report generated yet.", 404
        
    with open(TEMP_SUMMARY_PATH, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    doc = Document()
    doc.add_heading('AI Strategic Marketing Report', 0)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.find_all(['h3', 'p', 'li']):
        if element.name == 'h3':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'li':
            doc.add_paragraph(element.get_text(), style='List Bullet')
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
            
    word_path = os.path.join(DATA_DIR, "marketing_report.docx")
    doc.save(word_path)
    return send_file(word_path, as_attachment=True, download_name="AI_Insight_Report.docx")

@app.route('/validate')
def validate():
    gold_path = os.path.join(DATA_DIR, 'gold_standard.json')
    if not os.path.exists(gold_path): return "Gold standard file missing", 404
    
    dual_mode = request.args.get('dual_mode') == 'on'
    # Batch size parameter removed here; validator reads from config automatically
    metrics = run_gold_standard_validation(gold_path, dual_mode=dual_mode)
    return render_template('validation.html', metrics=metrics)

# Helper function: Compare and merge two runs
def compare_and_merge_runs(df1, df2):
    df1 = df1.reset_index(drop=True)
    df2 = df2.reset_index(drop=True)
    total = len(df1)
    consistent = 0
    df1['verification_status'] = 'Pending'
    for i in range(total):
        if i >= len(df2): break
        r1, r2 = df1.iloc[i], df2.iloc[i]
        if r1['aspect'] == r2['aspect'] and r1['sentiment'] == r2['sentiment']:
            consistent += 1
            df1.at[i, 'verification_status'] = '✅ Verified'
        else:
            df1.at[i, 'verification_status'] = f"⚠️ Mismatch (Run2: {r2['aspect']}/{r2['sentiment']})"
    score = (consistent / total * 100) if total > 0 else 0
    return df1, score

if __name__ == '__main__':
    app.run(debug=True, port=5000, threaded=True)